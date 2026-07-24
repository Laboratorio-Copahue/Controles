import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from reporte_farmacity.traductor import buscar_productos_en_texto


def set_cell_background(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def set_cell_widths(row, widths):
    for i, width in enumerate(widths):
        row.cells[i].width = width


def generar_reportes(df, col_comentario="COMENTARIOS", archivo_excel_salida="reporte_farmacity.xlsx", archivo_word_salida="reporte_farmacity.docx"):
    """
    Toma un DataFrame cargado desde Streamlit y genera:
    1. Un documento Word
    2. Un archivo Excel
    """
    
    # Aseguramos nombres de columnas estándar
    col_contacto = "CONTACTO" if "CONTACTO" in df.columns else df.columns[7]
    col_fecha = "FECHA" if "FECHA" in df.columns else df.columns[4]
    col_domicilio = "DOMICILIO" if "DOMICILIO" in df.columns else df.columns[10]
    col_comentarios = col_comentario if col_comentario in df.columns else "COMENTARIOS"

    # ==========================================
    # 1. GENERAR DOCUMENTO WORD
    # ==========================================
    doc = Document()
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(20)
    run_t = titulo.add_run("FALTANTES Y VISITAS FARMACITY: REPORTE SEMANAL")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(14)
    run_t.font.bold = True

    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = 'Table Grid'
    col_widths = [Inches(1.8), Inches(0.9), Inches(1.8), Inches(3.0)]

    hdr = tabla.rows[0].cells
    hdr_titles = ['CONTACTO', 'FECHA', 'DOMICILIO', 'COMENTARIOS']
    for i, title in enumerate(hdr_titles):
        hdr[i].text = title
        set_cell_background(hdr[i], "F2F2F2")
        if hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].font.bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    set_cell_widths(tabla.rows[0], col_widths)

    rows_excel = []

    for _, fila in df.iterrows():
        suc = str(fila[col_contacto]).strip() if pd.notna(fila[col_contacto]) else ""
        fecha = str(fila[col_fecha]).split()[0] if pd.notna(fila[col_fecha]) else ""
        domicilio = str(fila[col_domicilio]).strip() if pd.notna(fila[col_domicilio]) else ""
        comentario_orig = str(fila[col_comentarios]).strip() if pd.notna(fila[col_comentarios]) else ""

        if not suc or suc.lower() == "nan":
            continue

        row_cells = tabla.add_row().cells
        set_cell_widths(tabla.rows[-1], col_widths)

        row_cells[0].text = suc
        row_cells[1].text = fecha
        row_cells[2].text = domicilio

        celda_comentario = row_cells[3]
        
        # Buscar productos en el texto del comentario
        productos = buscar_productos_en_texto(comentario_orig)

        if productos and "acción en punto de venta" not in comentario_orig.lower():
            p_int = celda_comentario.paragraphs[0]
            run_f = p_int.add_run("FALTANTES:\n")
            run_f.bold = True
            run_f.font.size = Pt(9)

            # Subtabla de productos
            sub_t = celda_comentario.add_table(rows=1, cols=3)
            sub_t.style = 'Table Grid'
            sub_widths = [Inches(0.6), Inches(1.0), Inches(1.4)]

            for i, h in enumerate(['CUF', 'EAN', 'DESC']):
                sub_t.rows[0].cells[i].text = h
                set_cell_background(sub_t.rows[0].cells[i], "F9F9F9")
                if sub_t.rows[0].cells[i].paragraphs[0].runs:
                    sub_t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                    sub_t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(8)

            set_cell_widths(sub_t.rows[0], sub_widths)

            for prod in productos:
                sub_r = sub_t.add_row().cells
                set_cell_widths(sub_t.rows[-1], sub_widths)
                sub_r[0].text = str(prod["cuf"])
                sub_r[1].text = str(prod["ean"])
                sub_r[2].text = str(prod["desc"])

                for c in sub_r:
                    if c.paragraphs[0].runs:
                        c.paragraphs[0].runs[0].font.size = Pt(8)

                rows_excel.append({
                    "SUCURSAL": suc,
                    "FECHA": fecha,
                    "DIRECCIÓN": domicilio,
                    "CUF": str(prod["cuf"]),
                    "PRODUCTO": str(prod["desc"]),
                    "EAN": str(prod["ean"])
                })

            if comentario_orig and comentario_orig.lower() != "nan":
                p_obs = celda_comentario.add_paragraph()
                p_obs.paragraph_format.space_before = Pt(4)
                run_obs = p_obs.add_run(f"Obs: {comentario_orig}")
                run_obs.font.size = Pt(8)
                run_obs.font.italic = True
        else:
            celda_comentario.text = comentario_orig if comentario_orig and comentario_orig.lower() != "nan" else "Visita"

        for cell in row_cells[:3]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.save(archivo_word_salida)

    # ==========================================
    # 2. GENERAR ARCHIVO EXCEL
    # ==========================================
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Resumen Faltantes"
    ws.views.sheetView[0].showGridLines = True

    headers = ["SUCURSAL", "FECHA", "DIRECCIÓN", "CUF", "PRODUCTO", "EAN"]
    ws.append(headers)

    for r in rows_excel:
        ws.append([r["SUCURSAL"], r["FECHA"], r["DIRECCIÓN"], r["CUF"], r["PRODUCTO"], r["EAN"]])

    # Aplicar estilos básicos
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
    
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=6):
        for cell in row:
            cell.font = Font(name="Arial", size=9)
            cell.number_format = '@'

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

    wb.save(archivo_excel_salida)

    # Retornamos los DataFrames resultantes si los querés mostrar en pantalla en Streamlit
    df_excel = pd.DataFrame(rows_excel)
    return df_excel, df